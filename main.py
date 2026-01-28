from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel, Field
from typing import List, Optional
import csv
import os
from datetime import datetime
import io
import pandas as pd
from fpdf import FPDF
from docx import Document
from motor.motor_asyncio import AsyncIOMotorClient
import cloudinary
import cloudinary.uploader
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = FastAPI()

# Environment Validation
MONGODB_URI = os.getenv("MONGODB_URI")
CLOUDINARY_NAME = os.getenv("CLOUDINARY_CLOUD_NAME")
CLOUDINARY_KEY = os.getenv("CLOUDINARY_API_KEY")
CLOUDINARY_SECRET = os.getenv("CLOUDINARY_API_SECRET")

# Check if we are running in production (Render usually sets RENDER=true or similar)
IS_PRODUCTION = os.getenv("RENDER") == "true" or os.getenv("PORT") is not None

if not MONGODB_URI:
    if IS_PRODUCTION:
        raise RuntimeError("CRITICAL ERROR: MONGODB_URI is not set in environment variables.")
    else:
        print("WARNING: MONGODB_URI not found, defaulting to localhost.")
        MONGODB_URI = "mongodb://localhost:27017"

# MongoDB Configuration
client = AsyncIOMotorClient(MONGODB_URI)
db = client.get_database("inventory_db")
medicines_collection = db.get_collection("medicines")
users_collection = db.get_collection("users")
activities_collection = db.get_collection("activities")

# Cloudinary Configuration
if not all([CLOUDINARY_NAME, CLOUDINARY_KEY, CLOUDINARY_SECRET]):
    print("WARNING: Cloudinary credentials missing. Profile picture uploads will fail.")
else:
    cloudinary.config(
        cloud_name=CLOUDINARY_NAME,
        api_key=CLOUDINARY_KEY,
        api_secret=CLOUDINARY_SECRET,
        secure=True
    )

origins = [
    "http://localhost:3000",
    "https://medicine-inventory-management-syste-tau.vercel.app",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Files
MEDICINE_FILE = "medicines.txt"
USER_FILE = "user.txt"
LOG_FILE = "log.txt"

# Models
class Medicine(BaseModel):
    id: int
    name: str
    quantity: int
    price: float
    category: str
    manufacturer: str
    batchNumber: str
    minStock: int
    expiryDate: str
    status: str
    supplier: Optional[str] = None # Username or Name of supplier

class User(BaseModel):
    username: str
    email: str
    password: str
    role: str
    # Personal info
    firstName: str = ""
    lastName: str = ""
    gender: str = ""
    phoneNumber: str = ""
    profilePic: str = "" # URL or Base64
    address: str = "" # Primarily for suppliers

class UserCreate(User):
    pass

class UserPublic(BaseModel):
    username: str
    email: str
    role: str
    firstName: str = ""
    lastName: str = ""

class UserProfile(BaseModel):
    username: str
    email: str
    role: str
    firstName: str
    lastName: str
    gender: str
    phoneNumber: str
    profilePic: str
    address: str

class UserProfileUpdate(BaseModel):
    firstName: Optional[str] = None
    lastName: Optional[str] = None
    gender: Optional[str] = None
    phoneNumber: Optional[str] = None
    profilePic: Optional[str] = None
    address: Optional[str] = None

class UserLoginUpdate(BaseModel):
    username: Optional[str] = None
    email: Optional[str] = None
    newPassword: Optional[str] = None
    currentPassword: str

class MedicineCreate(BaseModel):
    name: str
    category: str
    manufacturer: str
    batchNumber: str
    quantity: int
    expiryDate: str
    price: float
    minStock: int
    supplier: Optional[str] = None

class Activity(BaseModel):
    timestamp: str
    user: str
    action: str
    details: str

class Alert(BaseModel):
    medicineId: int
    medicineName: str
    currentStock: int
    minStock: int
    status: str
    category: str
    manufacturer: str
    batchNumber: str
    expiryDate: str
    price: float
    timestamp: str
    daysRemaining: Optional[int] = None


# Helpers
def get_timestamp():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def log_activity(activity: str):
    with open(LOG_FILE, "a") as f:
        f.write(f"[{get_timestamp()}] {activity}\n")

async def log_activity_to_db(user: str, action: str, details: str):
    """Log activity to MongoDB for tracking"""
    activity = Activity(
        timestamp=get_timestamp(),
        user=user,
        action=action,
        details=details
    )
    await activities_collection.insert_one(activity.dict())

def get_stock_status(quantity: int, min_stock: int = 30, expiry_date: str = "N/A") -> str:
    # Check Expiry first as it's a critical safety concern
    if expiry_date and expiry_date != "N/A":
        try:
            now = datetime.now()
            expiry = datetime.strptime(expiry_date, "%Y-%m-%d")
            days_remaining = (expiry - now).days
            if days_remaining <= 90:
                return "Expiring Soon"
        except ValueError:
            pass

    # Use min_stock as threshold for Low Stock, half of it for Critical
    # If min_stock is 0, default to current hardcoded values
    low_threshold = min_stock if min_stock > 0 else 30
    critical_threshold = low_threshold // 2 if low_threshold > 0 else 15
    
    if quantity < critical_threshold:
        return "Critical"
    elif quantity < low_threshold:
        return "Low Stock"
    else:
        return "In Stock"

# Helper for MongoDB Migration
async def migrate_data():
    # Migrate Medicines
    if await medicines_collection.count_documents({}) == 0 and os.path.exists(MEDICINE_FILE):
        print("Migrating medicines from file to MongoDB...")
        # Use existing sync load_medicines logic for one-time migration
        with open(MEDICINE_FILE, "r", newline='', encoding='utf-8') as f:
            first_line = f.readline()
            f.seek(0)
            if first_line:
                has_header = first_line.startswith("id,")
                meds_to_insert = []
                if has_header:
                    reader = csv.DictReader(f)
                    for row in reader:
                        try:
                            meds_to_insert.append({
                                "id": int(row["id"]),
                                "name": row["name"],
                                "quantity": int(row["quantity"]),
                                "price": float(row["price"]),
                                "category": row["category"],
                                "manufacturer": row["manufacturer"],
                                "batchNumber": row["batchNumber"],
                                "minStock": int(row["minStock"]),
                                "expiryDate": row["expiryDate"],
                                "expiryDate": row["expiryDate"],
                                "status": row["status"],
                                "supplier": row.get("supplier", None)
                            })
                        except: continue
                if meds_to_insert:
                    await medicines_collection.insert_many(meds_to_insert)

    # Migrate Users
    if await users_collection.count_documents({}) == 0 and os.path.exists(USER_FILE):
        print("Migrating users from file to MongoDB...")
        with open(USER_FILE, "r", newline='', encoding='utf-8') as f:
            first_line = f.readline()
            f.seek(0)
            if first_line:
                has_header = first_line.startswith("username,")
                users_to_insert = []
                if has_header:
                    reader = csv.DictReader(f)
                    for row in reader:
                        users_to_insert.append(row)
                else:
                    reader = csv.reader(f)
                    for parts in reader:
                        if len(parts) == 4:
                            users_to_insert.append({
                                "username": parts[0],
                                "email": parts[1],
                                "password": parts[2],
                                "role": parts[3],
                                "firstName": "", "lastName": "", "gender": "",
                                "phoneNumber": "", "profilePic": "", "address": ""
                            })
                if users_to_insert:
                    await users_collection.insert_many(users_to_insert)

@app.on_event("startup")
async def startup_db_client():
    await migrate_data()

# Data Access functions (now async)
async def load_medicines() -> List[Medicine]:
    meds = []
    cursor = medicines_collection.find({})
    async for doc in cursor:
        doc.pop('_id', None)
        meds.append(Medicine(**doc))
    return meds

async def save_medicines(meds: List[Medicine]):
    # In MongoDB, we usually update specific documents
    # But for compatibility with existing "overwrite all" logic:
    for med in meds:
        med.status = get_stock_status(med.quantity, med.minStock, med.expiryDate)
        await medicines_collection.replace_one({"id": med.id}, med.dict(), upsert=True)

async def load_users() -> List[User]:
    users = []
    cursor = users_collection.find({})
    async for doc in cursor:
        doc.pop('_id', None)
        users.append(User(**doc))
    return users

async def save_users(users: List[User]):
    for user in users:
        await users_collection.replace_one({"username": user.username}, user.dict(), upsert=True)

# Security
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

async def get_current_user(token: str = Depends(oauth2_scheme)):
    users = await load_users()
    for user in users:
        if user.username == token:
            return user
    raise HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Invalid authentication credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )

def require_role(allowed_roles: List[str]):
    def role_checker(current_user: User = Depends(get_current_user)):
        if current_user.role not in allowed_roles:
            raise HTTPException(status_code=403, detail="Operation not permitted for this user role")
        return current_user
    return role_checker

# Endpoints

@app.post("/token")
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    users = await load_users()
    for user in users:
        if (user.username == form_data.username or user.email == form_data.username) and user.password == form_data.password:
            log_activity(f"User {user.username} logged in as {user.role}")
            await log_activity_to_db(user.username, "LOGIN", f"User logged in as {user.role}")
            return {"access_token": user.username, "token_type": "bearer"}
    raise HTTPException(status_code=400, detail="Incorrect username, email or password")

@app.get("/users/me", response_model=UserPublic)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return UserPublic(**current_user.dict())

@app.get("/users/me/profile", response_model=UserProfile)
async def get_my_profile(current_user: User = Depends(get_current_user)):
    return UserProfile(**current_user.dict())

@app.put("/users/me/profile", response_model=UserProfile)
async def update_my_profile(profile_update: UserProfileUpdate, current_user: User = Depends(get_current_user)):
    users = await load_users()
    for i, u in enumerate(users):
        if u.username == current_user.username:
            updated_data = u.dict()
            # Only update provided fields
            for field, value in profile_update.dict(exclude_unset=True).items():
                if value is not None:
                    updated_data[field] = value
            
            updated_user = User(**updated_data)
            users[i] = updated_user
            await save_users(users)
            log_activity(f"User {u.username} updated profile info")
            await log_activity_to_db(current_user.username, "UPDATE_PROFILE", f"Updated personal information")
            return UserProfile(**updated_user.dict())
    raise HTTPException(status_code=404, detail="User not found")

@app.put("/users/me/login-details")
async def update_login_details(login_update: UserLoginUpdate, current_user: User = Depends(get_current_user)):
    # Verify current password
    if login_update.currentPassword != current_user.password:
        raise HTTPException(status_code=400, detail="Incorrect current password")
    
    users = await load_users()
    for i, u in enumerate(users):
        if u.username == current_user.username:
            # Check username uniqueness if changing
            if login_update.username and login_update.username != current_user.username:
                if any(other.username == login_update.username for other in users):
                    raise HTTPException(status_code=400, detail="Username already taken")
                u.username = login_update.username
            
            # Check email uniqueness if changing
            if login_update.email and login_update.email != current_user.email:
                if any(other.email == login_update.email for other in users):
                    raise HTTPException(status_code=400, detail="Email already registered")
                u.email = login_update.email
            
            # Update password if provided
            if login_update.newPassword:
                u.password = login_update.newPassword
            
            users[i] = u
            await save_users(users)
            log_activity(f"User {current_user.username} updated login details")
            await log_activity_to_db(current_user.username, "UPDATE_LOGIN", f"Updated login credentials")
            return {"message": "Login details updated successfully. Please re-login if username was changed."}
            
    raise HTTPException(status_code=404, detail="User not found")

@app.get("/medicines", response_model=List[Medicine])
async def get_medicines(current_user: User = Depends(get_current_user)):
    return await load_medicines()

@app.get("/medicines/search")
async def search_medicine(name: str, current_user: User = Depends(get_current_user)):
    meds = await load_medicines()
    # Case insensitive substring search
    return [m for m in meds if name.lower() in m.name.lower()]

@app.get("/medicines/status")
async def get_medicines_status(current_user: User = Depends(get_current_user)):
    """
    Returns categorization of medicines logic.
    User asked for endpoint to measure quantity:
    < 30 -> Low Stock
    < 15 -> Critical
    """
    meds = await load_medicines()
    status_report = {
        "Critical": [m for m in meds if m.status == "Critical"],
        "Low Stock": [m for m in meds if m.status == "Low Stock"],
        "Expiring Soon": [m for m in meds if m.status == "Expiring Soon"],
        "In Stock": [m for m in meds if m.status == "In Stock"]
    }
    return status_report

@app.get("/medicines/export")
async def export_medicines(format: str, current_user: User = Depends(get_current_user)):
    meds = await load_medicines()
    if not meds:
        raise HTTPException(status_code=404, detail="No medicines to export")
    
    # Convert pydantic models to list of dicts
    data = [m.dict() for m in meds]
    df = pd.DataFrame(data)
    
    if format.lower() == "csv":
        stream = io.StringIO()
        df.to_csv(stream, index=False)
        response = StreamingResponse(iter([stream.getvalue()]), media_type="text/csv")
        response.headers["Content-Disposition"] = "attachment; filename=medicines.csv"
        return response

    elif format.lower() == "pdf":
        pdf = FPDF(orientation='L', unit='mm', format='A4')
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, txt="Medicine Inventory Report", ln=1, align='C')
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(0, 10, txt=f"Generated on: {get_timestamp()}", ln=1, align='R')
        pdf.ln(5)
        
        # Table Header
        pdf.set_font("Arial", 'B', 9)
        # ID, Name, Qty, Price, Category, Manufacturer, Batch, Min Stock, Expiry, Status
        headers = ["ID", "Name", "Qty", "Price", "Category", "Manufacturer", "Batch #", "Min Stock", "Expiry", "Status"]
        # Total width for Landscape A4 is ~277mm (297 - 2*10 margin)
        widths = [10, 40, 15, 20, 30, 40, 30, 20, 30, 30] 
        
        for i, col in enumerate(headers):
            pdf.cell(widths[i], 10, col, 1, 0, 'C')
        pdf.ln()
        
        # Table Body
        pdf.set_font("Arial", size=8)
        for med in meds:
            pdf.cell(widths[0], 10, str(med.id), 1, 0, 'C')
            pdf.cell(widths[1], 10, str(med.name)[:20], 1)
            pdf.cell(widths[2], 10, str(med.quantity), 1, 0, 'C')
            pdf.cell(widths[3], 10, f"${med.price:.2f}", 1, 0, 'R')
            pdf.cell(widths[4], 10, str(med.category)[:15], 1)
            pdf.cell(widths[5], 10, str(med.manufacturer)[:20], 1)
            pdf.cell(widths[6], 10, str(med.batchNumber), 1)
            pdf.cell(widths[7], 10, str(med.minStock), 1, 0, 'C')
            pdf.cell(widths[8], 10, str(med.expiryDate), 1, 0, 'C')
            pdf.cell(widths[9], 10, str(med.status), 1, 0, 'C')
            pdf.ln()
            
        stream = io.BytesIO()
        pdf_content = pdf.output(dest='S').encode('latin-1')
        stream.write(pdf_content)
        stream.seek(0)
        
        response = StreamingResponse(stream, media_type="application/pdf")
        response.headers["Content-Disposition"] = "attachment; filename=medicines.pdf"
        return response

    elif format.lower() == "docx":
        doc = Document()
        doc.add_heading('Medicine Inventory Report', 0)
        doc.add_paragraph(f"Generated on: {get_timestamp()}")
        
        # ID, Name, Qty, Price, Category, Manufacturer, Batch, Min Stock, Expiry, Status
        table = doc.add_table(rows=1, cols=10)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["ID", "Name", "Qty", "Price", "Category", "Manufacturer", "Batch #", "Min Stock", "Expiry", "Status"]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for med in meds:
            row_cells = table.add_row().cells
            row_cells[0].text = str(med.id)
            row_cells[1].text = med.name
            row_cells[2].text = str(med.quantity)
            row_cells[3].text = f"${med.price:.2f}"
            row_cells[4].text = med.category
            row_cells[5].text = med.manufacturer
            row_cells[6].text = med.batchNumber
            row_cells[7].text = str(med.minStock)
            row_cells[8].text = med.expiryDate
            row_cells[9].text = med.status

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        response = StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response.headers["Content-Disposition"] = "attachment; filename=medicines.docx"
        return response
        
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Supported: csv, pdf, docx")

@app.post("/medicines", dependencies=[Depends(require_role(["admin"]))], response_model=Medicine)
async def add_medicine(med: MedicineCreate, current_user: User = Depends(get_current_user)):
    if med.quantity < 0:
        raise HTTPException(status_code=400, detail="Quantity cannot be negative")
    if med.price < 0:
        raise HTTPException(status_code=400, detail="Price cannot be negative")

    meds = await load_medicines()
    next_id = 1
    if meds:
        next_id = max(m.id for m in meds) + 1
    
    stat = get_stock_status(med.quantity, med.minStock, med.expiryDate)
    
    new_med = Medicine(
        id=next_id, 
        name=med.name, 
        quantity=med.quantity, 
        price=med.price,
        category=med.category,
        manufacturer=med.manufacturer,
        batchNumber=med.batchNumber,
        minStock=med.minStock,
        expiryDate=med.expiryDate,
        status=stat,
        supplier=med.supplier
    )
    meds.append(new_med)
    await save_medicines(meds)
    log_activity(f"Added medicine: {med.name} (Supplier: {med.supplier})")
    return new_med

@app.put("/medicines/{med_id}", dependencies=[Depends(require_role(["admin"]))], response_model=Medicine)
async def update_medicine(med_id: int, med_update: MedicineCreate, current_user: User = Depends(get_current_user)):
    if med_update.quantity < 0:
        raise HTTPException(status_code=400, detail="Quantity cannot be negative")
    if med_update.price < 0:
        raise HTTPException(status_code=400, detail="Price cannot be negative")
        
    meds = await load_medicines()
    for i, m in enumerate(meds):
        if m.id == med_id:
            # Update fields
            m.name = med_update.name
            m.quantity = med_update.quantity
            m.price = med_update.price
            m.category = med_update.category
            m.manufacturer = med_update.manufacturer
            m.batchNumber = med_update.batchNumber
            m.minStock = med_update.minStock
            m.expiryDate = med_update.expiryDate
            m.supplier = med_update.supplier
            m.status = get_stock_status(m.quantity, m.minStock, m.expiryDate) # Recalculate status
            
            meds[i] = m
            await save_medicines(meds)
            log_activity(f"Updated medicine ID {med_id}")
            return m
    raise HTTPException(status_code=404, detail="Medicine not found")

@app.delete("/medicines/{med_id}", dependencies=[Depends(require_role(["admin"]))])
async def delete_medicine(med_id: int, current_user: User = Depends(get_current_user)):
    meds = await load_medicines()
    initial_len = len(meds)
    meds = [m for m in meds if m.id != med_id]
    if len(meds) == initial_len:
         raise HTTPException(status_code=404, detail="Medicine not found")
    
    await save_medicines(meds)
    log_activity(f"Deleted medicine ID {med_id}")
    return {"message": "Medicine deleted successfully"}

@app.put("/medicines/{med_id}/restock", dependencies=[Depends(require_role(["admin", "supplier"]))])
async def restock_medicine(med_id: int, amount: int, current_user: User = Depends(get_current_user)):
    meds = await load_medicines()
    for m in meds:
        if m.id == med_id:
            m.quantity += amount
            m.status = get_stock_status(m.quantity, m.minStock, m.expiryDate)
            await save_medicines(meds)
            log_activity(f"Restocked medicine ID {med_id} by {amount}")
            return m
    raise HTTPException(status_code=404, detail="Medicine not found")

@app.put("/medicines/{med_id}/dispense", dependencies=[Depends(require_role(["admin", "pharmacist"]))])
async def dispense_medicine(med_id: int, amount: int, current_user: User = Depends(get_current_user)):
    meds = await load_medicines()
    for m in meds:
        if m.id == med_id:
            if m.quantity < amount:
                raise HTTPException(status_code=400, detail="Insufficient stock")
            m.quantity -= amount
            m.status = get_stock_status(m.quantity, m.minStock, m.expiryDate)
            await save_medicines(meds)
            log_activity(f"Dispensed medicine ID {med_id} amount {amount}")
            await log_activity_to_db(current_user.username, "DISPENSE_MEDICINE", f"Dispensed {amount} units of {m.name} (ID: {med_id})")
            return m
    raise HTTPException(status_code=404, detail="Medicine not found")

@app.get("/users", dependencies=[Depends(require_role(["admin"]))], response_model=List[UserProfile])
async def read_users(current_user: User = Depends(get_current_user)):
    users = await load_users()
    return [UserProfile(**u.dict()) for u in users]

@app.get("/suppliers", dependencies=[Depends(require_role(["admin", "pharmacist"]))], response_model=List[UserPublic])
async def get_suppliers(current_user: User = Depends(get_current_user)):
    users = await load_users()
    return [UserPublic(**u.dict()) for u in users if u.role == "supplier"]

@app.get("/users/{username}", dependencies=[Depends(require_role(["admin"]))], response_model=UserProfile)
async def read_user_detail(username: str, current_user: User = Depends(get_current_user)):
    users = await load_users()
    for u in users:
        if u.username == username:
            return UserProfile(**u.dict())
    raise HTTPException(status_code=404, detail="User not found")

@app.post("/users", dependencies=[Depends(require_role(["admin"]))])
async def create_user(user: User, current_user: User = Depends(get_current_user)):
    users = await load_users()
    if any(u.username == user.username for u in users):
        raise HTTPException(status_code=400, detail="Username already exists")
    if any(u.email == user.email for u in users):
        raise HTTPException(status_code=400, detail="Email already exists")
    
    # Auto-create login credentials: email as username, healme12 as default password
    user.username = user.email
    user.password = "healme12"
    
    users.append(user)
    await save_users(users)
    log_activity(f"Admin added new user: {user.username} ({user.role})")
    await log_activity_to_db(current_user.username, "CREATE_USER", f"Created user {user.username} with role {user.role}")
    return {"message": "User added successfully", "username": user.email, "defaultPassword": "healme12"}

@app.put("/users/{username}", dependencies=[Depends(require_role(["admin"]))])
async def update_any_user(username: str, user_update: User, current_user: User = Depends(get_current_user)):
    # Update user in MongoDB directly to avoid duplicates
    users = await load_users()
    user_found = None
    
    for u in users:
        if u.username == username:
            user_found = u
            break
    
    if not user_found:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Check username uniqueness if changing
    if user_update.username != username:
        if any(other.username == user_update.username for other in users):
            raise HTTPException(status_code=400, detail="New username already taken")
    
    # Update in MongoDB using replace_one with username filter
    await users_collection.replace_one({"username": username}, user_update.dict())
    log_activity(f"Admin updated user: {username}")
    await log_activity_to_db(current_user.username, "UPDATE_USER", f"Updated user {username}")
    return {"message": "User updated successfully"}

@app.delete("/users/{username}", dependencies=[Depends(require_role(["admin"]))])
async def delete_user(username: str, current_user: User = Depends(get_current_user)):
    # Check self-deletion BEFORE filtering
    if username == current_user.username:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    
    # Delete from MongoDB directly
    result = await users_collection.delete_one({"username": username})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    log_activity(f"Admin deleted user: {username}")
    await log_activity_to_db(current_user.username, "DELETE_USER", f"Deleted user {username}")
    return {"message": "User deleted successfully"}

from fastapi import File, UploadFile

@app.post("/users/me/profile-pic")
async def upload_profile_pic(profilePic: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    try:
        # Upload to Cloudinary
        result = cloudinary.uploader.upload(profilePic.file, folder="ims_profiles")
        url = result.get("secure_url")
        
        # Update user profile
        users = await load_users()
        for i, u in enumerate(users):
            if u.username == current_user.username:
                u.profilePic = url
                users[i] = u
                await save_users(users)
                log_activity(f"User {u.username} uploaded new profile picture")
                return {"profilePic": url}
        
        raise HTTPException(status_code=404, detail="User not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

# Activity Logging Endpoint
@app.get("/activities", dependencies=[Depends(require_role(["admin"]))], response_model=List[Activity])
async def get_activities(current_user: User = Depends(get_current_user), limit: int = 100):
    """Get recent activities (admin only)"""
    activities = []
    cursor = activities_collection.find({}).sort("timestamp", -1).limit(limit)
    async for doc in cursor:
        doc.pop('_id', None)
        activities.append(Activity(**doc))
    return activities

# Alert Endpoints

@app.get("/alerts", dependencies=[Depends(require_role(["admin", "pharmacist"]))], response_model=List[Alert])
async def get_alerts_enhanced(current_user: User = Depends(get_current_user)):
    """
    Get medicines with:
    1. Low Stock or Critical status
    2. Expiry date within 3 months (90 days)
    """
    meds = await load_medicines()
    alerts = []
    now = datetime.now()
    
    for med in meds:
        should_alert = False
        days_remaining = None
        
        # Check Stock or Expiry status
        if med.status in ["Low Stock", "Critical", "Expiring Soon"]:
            should_alert = True
            
        # Check Expiry details for daysRemaining field
        try:
            # Assuming YYYY-MM-DD
            expiry = datetime.strptime(med.expiryDate, "%Y-%m-%d")
            delta = expiry - now
            days_remaining = delta.days
            
            if days_remaining <= 90:
                should_alert = True
                if med.status == "In Stock": # Override status for display purpose if it's purely an expiry alert? 
                    # Actually, the user wants "in the alert it should show... days remaining"
                    # We can keep original status but maybe add a note or just rely on the client to show daysRemaining
                    pass
        except ValueError:
            # Invalid date format, skip expiry check or log warning
            pass

        if should_alert:
            alerts.append(Alert(
                medicineId=med.id,
                medicineName=med.name,
                currentStock=med.quantity,
                minStock=med.minStock,
                status=med.status, # Keep original stock status
                category=med.category,
                manufacturer=med.manufacturer,
                batchNumber=med.batchNumber,
                expiryDate=med.expiryDate,
                price=med.price,
                timestamp=get_timestamp(),
                daysRemaining=days_remaining
            ))
            
    return alerts

@app.get("/medicines/expiring", response_model=List[Medicine])
async def get_expiring_medicines(current_user: User = Depends(get_current_user)):
    """Fetch medicines that are expiring soon (within 90 days)"""
    meds = await load_medicines()
    return [m for m in meds if m.status == "Expiring Soon"]

@app.get("/supplier/alerts", dependencies=[Depends(require_role(["supplier"]))], response_model=List[Alert])
async def get_supplier_alerts(current_user: User = Depends(get_current_user)):
    """Get low stock alerts for medicines relevant to this supplier"""
    meds = await load_medicines()
    alerts = []
    
    # For now, show all low stock items to suppliers
    # In future, you could filter by manufacturer or add supplier assignment to medicines
    for med in meds:
        if med.status in ["Low Stock", "Critical"]:
            alerts.append(Alert(
                medicineId=med.id,
                medicineName=med.name,
                currentStock=med.quantity,
                minStock=med.minStock,
                status=med.status,
                category=med.category,
                manufacturer=med.manufacturer,
                batchNumber=med.batchNumber,
                expiryDate=med.expiryDate,
                price=med.price,
                timestamp=get_timestamp(),
                daysRemaining=None # Explicitly None or calc if needed
            ))
    return alerts
